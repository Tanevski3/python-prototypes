import docx
import re

location_to_file = "South Park_ Reel Chaos - FINAL Math.docx"
read_doc = docx.Document(location_to_file)

def match_text(text, para):
    match = re.match("^" + text + "\s*?$", para.text)
    print(match)
    return match

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    count = 0
    for para in doc.paragraphs:
        count += 1
        fullText.append( str(count) + " " + para.text + " " + para.style.name)
    return '\n'.join(fullText)


print(getText(location_to_file))

